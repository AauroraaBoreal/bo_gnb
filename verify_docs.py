# Automated Test: verify_docs.py
# Verifies that DOCX templating and PDF generation workflows run successfully without exceptions

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def test_docx_generation(template_path, output_path):
    print("Testing DOCX generation...")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")
        
    doc = Document(template_path)
    
    # 1. Test paragraphs replace
    replacements = {
        "{{quotation_date}}": "Callao, 02 de Julio del 2026",
        "{{client_name}}": "PRODAC SA",
        "{{client_attention}}": "Ing. Teodoro Tacsa",
        "{{quotation_number}}": "COTIZACION N° 999",
        "{{igv_text}}": "Precio no Incluye IGV.",
        "{{currency_text}}": "Moneda: Soles",
        "{{total_amount}}": "S/ 1,045.00",
        "{{total_in_words}}": "(Un mil cuarenta y cinco con 00/100 Soles)"
    }
    
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
                
    # 2. Test table row expansion and merging
    assert len(doc.tables) > 0, "No tables found in template!"
    table = doc.tables[0]
    
    # Clean template placeholder rows (keep header row 1, row 0 is empty)
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)
        
    # Write mock items (Item 1 with 2 rows description)
    # Check row capacity and add if necessary
    for r_idx in range(2, 4):
        if r_idx >= len(table.rows):
            table.add_row()
            
    # Row 2 (first item line)
    row2 = table.rows[2]
    row2.cells[0].text = "1"
    row2.cells[1].text = "Servicio de Reparación de 11 Spider en mal estado , soldando con supercito de ⅛."
    row2.cells[2].text = "11 Und"
    row2.cells[3].text = "95.00"
    row2.cells[4].text = "1,045.00"
    
    # Row 3 (second item description line)
    row3 = table.rows[3]
    row3.cells[0].text = "1"
    row3.cells[1].text = "Enderezando las partes dobladas con equipo oxicorte"
    row3.cells[2].text = "11 Und"
    row3.cells[3].text = "95.00"
    row3.cells[4].text = "1,045.00"
    
    # Merge cells vertically in columns 0, 2, 3, and 4
    for col in [0, 2, 3, 4]:
        first_cell = table.cell(2, col)
        cell_to_merge = table.cell(3, col)
        first_cell.merge(cell_to_merge)
        
    doc.save(output_path)
    print(f"DOCX test passed. File saved to: {output_path}")

def test_pdf_generation(output_path):
    print("Testing PDF generation...")
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TestTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1,
        textColor=colors.HexColor('#1E3D59')
    )
    
    story.append(Paragraph("GNB SOLUCIONES INDUSTRIALES S.A.C", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Este es un PDF de prueba generado para verificar el funcionamiento de ReportLab.", styles['Normal']))
    
    doc.build(story)
    print(f"PDF test passed. File saved to: {output_path}")

if __name__ == "__main__":
    template = r"C:\Users\Dayana\.gemini\antigravity\scratch\gnb_backoffice\templates\cotizacion_base.docx"
    
    out_docx = r"C:\Users\Dayana\.gemini\antigravity\scratch\gnb_backoffice\verify_cotiz_test.docx"
    out_pdf = r"C:\Users\Dayana\.gemini\antigravity\scratch\gnb_backoffice\verify_cotiz_test.pdf"
    
    test_docx_generation(template, out_docx)
    test_pdf_generation(out_pdf)
    print("\n[OK] DOCUMENT COMPILER TESTS COMPLETED SUCCESSFULLY!")

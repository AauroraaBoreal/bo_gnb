import os
import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from lib.supabase_client import get_supabase_client
from lib.db import get_setting_value
from lib.utils import format_currency

# --- STORAGE SERVICE ---

def upload_document_to_supabase_storage(file_path: str, bucket: str, remote_path: str) -> str:
    """
    Uploads a local file to a Supabase Storage bucket.
    Returns the public URL of the uploaded file.
    """
    supabase = get_supabase_client()
    
    # Read file bytes
    with open(file_path, "rb") as f:
        file_data = f.read()
        
    # Check if bucket exists, if upload fails we catch it
    try:
        # Check if file already exists. If so, update it, otherwise upload it.
        # We can try uploading and catch the error, or use remove then upload.
        try:
            supabase.storage.from_(bucket).remove([remote_path])
        except Exception:
            pass
            
        supabase.storage.from_(bucket).upload(
            path=remote_path,
            file=file_data,
            file_options={"content-type": "application/octet-stream"}
        )
        
        # Get public URL
        url_res = supabase.storage.from_(bucket).get_public_url(remote_path)
        return url_res
    except Exception as e:
        raise IOError(f"Error al subir archivo a Supabase Storage ({bucket}/{remote_path}): {str(e)}")


def check_custom_template_exists() -> bool:
    """
    Checks if a custom template has been uploaded to Supabase Storage.
    """
    try:
        supabase = get_supabase_client()
        files = supabase.storage.from_("quotations").list()
        if files and isinstance(files, list):
            return any(isinstance(f, dict) and f.get("name") == "cotizacion_template.docx" for f in files)
        return False
    except Exception:
        return False


def get_template_file_path() -> str:
    """
    Returns the path to the template file.
    If a custom template exists in Supabase Storage ('cotizacion_template.docx'),
    it downloads it to a temporary location and returns that path.
    Otherwise, it returns the local template file path.
    """
    local_path = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates", "cotizacion_base.docx"))
    
    try:
        if check_custom_template_exists():
            supabase = get_supabase_client()
            import tempfile
            res = supabase.storage.from_("quotations").download("cotizacion_template.docx")
            tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp_file.write(res)
            tmp_file.close()
            return tmp_file.name
    except Exception as e:
        print(f"Error checking/downloading custom template: {e}")
        
    return local_path


def get_watermark_image_path() -> str:
    """
    Checks if a custom template exists, and if so, extracts its watermark image (word/media/image1.jpg)
    to a temporary file and returns its path.
    """
    try:
        template_path = get_template_file_path()
        local_base_path = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates", "cotizacion_base.docx"))
        
        # If the template path points to a downloaded custom template, we extract the image
        if template_path != local_base_path and os.path.exists(template_path):
            import zipfile
            import tempfile
            with zipfile.ZipFile(template_path, 'r') as z:
                # Find any image in word/media/
                media_files = [name for name in z.namelist() if name.startswith("word/media/")]
                if media_files:
                    # We pick the first image (usually image1.jpg or image1.png, which is the watermark)
                    img_data = z.read(media_files[0])
                    ext = media_files[0].split(".")[-1]
                    
                    # Write to a persistent temp file
                    tmp_dir = tempfile.gettempdir()
                    watermark_temp_path = os.path.join(tmp_dir, f"gnb_watermark.{ext}")
                    with open(watermark_temp_path, "wb") as f:
                        f.write(img_data)
                    return watermark_temp_path
    except Exception as e:
        print(f"Error extracting watermark image: {e}")
    return None


# --- DOCX GENERATION ---


def generate_docx_from_template(quotation_id: str, template_path: str, output_path: str):
    """
    Fills a DOCX template with quotation details.
    Dynamically expands the items table, vertically merging cells where appropriate.
    """
    supabase = get_supabase_client()
    
    # 1. Fetch quotation and client details
    quote = supabase.table("quotations").select("*, clients(*)").eq("id", quotation_id).execute().data[0]
    items = supabase.table("quotation_items").select("*").eq("quotation_id", quotation_id).order("item_order").execute().data
    
    # Prepare text placeholders
    q_date = datetime.datetime.strptime(quote["quotation_date"], "%Y-%m-%d")
    date_str = f"Callao, {q_date.day:02d} de {get_month_name(q_date.month)} del {q_date.year}"
    
    client_name = quote["clients"]["business_name"]
    client_attention = quote["attention_to"] or ""
    q_number = f"COTIZACION N° {quote['quotation_number']}"
    
    currency_label = "Soles" if quote["currency"] == "soles" else "Dólares"
    igv_label = "Precio no Incluye IGV." if not quote["include_igv"] else "Precio Incluye IGV."
    
    total_val = float(quote["total"])
    total_str = f"{'S/' if quote['currency'] == 'soles' else '$'} {total_val:,.2f}"
    total_words = f"({quote['total_in_words']})"
    
    replacements = {
        "{{quotation_date}}": date_str,
        "{{client_name}}": client_name,
        "{{client_attention}}": client_attention,
        "{{quotation_number}}": q_number,
        "{{igv_text}}": igv_label,
        "{{currency_text}}": f"Moneda: {currency_label}",
        "{{total_amount}}": total_str,
        "{{total_in_words}}": total_words
    }
    
    # Check if custom template has tables
    local_base_path = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates", "cotizacion_base.docx"))
    
    # 2. Open document.
    # If the custom template has no tables, we use the local base template as the layout,
    # and we will later merge the watermark header from the custom template.
    use_watermark_merge = False
    actual_template_path = template_path
    
    try:
        temp_doc = Document(template_path)
        has_tables = len(temp_doc.tables) > 0
        
        # If it doesn't look like a complete template, we use the base template and merge headers
        if not has_tables and template_path != local_base_path:
            actual_template_path = local_base_path
            use_watermark_merge = True
    except Exception as e:
        print(f"Error checking template structure, falling back to base template: {e}")
        actual_template_path = local_base_path
        use_watermark_merge = True
        
    doc = Document(actual_template_path)
    
    # Helper function to replace in a list of paragraphs
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, value)
                    
    # Helper to scan and replace in tables
    def replace_in_table(t):
        for row in t.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Replace in body paragraphs
    replace_in_paragraphs(doc.paragraphs)
    
    # Find the items table by checking for headers keywords
    items_table = None
    for t in doc.tables:
        found = False
        # Check first 5 rows for keywords
        for r_idx in range(min(5, len(t.rows))):
            try:
                row_texts = [cell.text.upper() for cell in t.rows[r_idx].cells]
                if any("SERVICIO" in txt or "ITEM" in txt or "DESCRIPCIÓN" in txt or "DESCRIPCION" in txt or "CANT" in txt for txt in row_texts):
                    items_table = t
                    found = True
                    break
            except Exception:
                pass
        if found:
            break
            
    if not items_table and len(doc.tables) > 0:
        items_table = doc.tables[0]
        
    # Replace in all other tables
    for t in doc.tables:
        if t is not items_table:
            replace_in_table(t)
            
    # Replace in headers and footers
    for section in doc.sections:
        if section.header:
            replace_in_paragraphs(section.header.paragraphs)
            for t in section.header.tables:
                replace_in_table(t)
        if section.footer:
            replace_in_paragraphs(section.footer.paragraphs)
            for t in section.footer.tables:
                replace_in_table(t)
                
    # 3. Populate items table
    if items_table:
        # Find header index and template row index in items table
        header_idx = 1
        template_idx = 2
        for r_idx in range(len(items_table.rows)):
            try:
                row_texts = [cell.text.upper() for cell in items_table.rows[r_idx].cells]
                if any("SERVICIO" in txt or "DESCRIPCIÓN" in txt or "DESCRIPCION" in txt for txt in row_texts):
                    header_idx = r_idx
                    template_idx = r_idx + 1 if r_idx + 1 < len(items_table.rows) else r_idx
                    break
            except Exception:
                pass
                
        # Clean extra rows after template_idx
        while len(items_table.rows) > (template_idx + 1):
            items_table._tbl.remove(items_table.rows[-1]._tr)
            
        # Ensure we have at least the template row
        while len(items_table.rows) < (template_idx + 1):
            items_table.add_row()
            
        row_idx = template_idx
        
        for item in items:
            desc = item["service_description"]
            # Split lines for multi-line description support
            lines = [line.strip() for line in desc.split("\n") if line.strip()]
            if not lines:
                lines = [""]
                
            start_row = row_idx
            
            for line in lines:
                if row_idx >= len(items_table.rows):
                    items_table.add_row()
                    
                row = items_table.rows[row_idx]
                
                # Safe fill of cell contents
                if len(row.cells) >= 5:
                    row.cells[0].text = str(item["item_order"])
                    row.cells[1].text = line
                    row.cells[2].text = f"{int(item['quantity'])} {item['unit']}"
                    row.cells[3].text = f"{float(item['unit_price']):,.2f}"
                    row.cells[4].text = f"{float(item['total']):,.2f}"
                elif len(row.cells) > 1:
                    row.cells[0].text = str(item["item_order"])
                    row.cells[1].text = line
                    
                row_idx += 1
                
            end_row = row_idx - 1
            
            # Merge cells vertically in columns 0, 2, 3, and 4
            if end_row > start_row:
                for col in [0, 2, 3, 4]:
                    if col < len(items_table.columns):
                        first_cell = items_table.cell(start_row, col)
                        for r in range(start_row + 1, end_row + 1):
                            cell = items_table.cell(r, col)
                            first_cell.merge(cell)
                            
    # Save the populated document
    doc.save(output_path)
    
    # 4. If we need to merge the watermark header from the custom template
    if use_watermark_merge:
        try:
            import zipfile
            import shutil
            
            # Extract header and media from the custom template
            watermark_files = {}
            with zipfile.ZipFile(template_path, 'r') as z_custom:
                for name in z_custom.namelist():
                    if name == "word/header1.xml" or name == "word/_rels/header1.xml.rels" or name.startswith("word/media/"):
                        watermark_files[name] = z_custom.read(name)
                        
            if watermark_files:
                temp_zip = output_path + ".tmp"
                with zipfile.ZipFile(output_path, 'r') as z_in:
                    with zipfile.ZipFile(temp_zip, 'w', zipfile.ZIP_DEFLATED) as z_out:
                        # Copy all files from populated docx except the ones we want to overwrite
                        for item in z_in.infolist():
                            if item.filename not in watermark_files:
                                z_out.writestr(item, z_in.read(item.filename))
                                
                        # Insert watermark files
                        for name, data in watermark_files.items():
                            z_out.writestr(name, data)
                
                shutil.move(temp_zip, output_path)
        except Exception as merge_err:
            print(f"Error merging watermark header: {merge_err}")


# --- PDF GENERATION ---

def generate_pdf_from_quotation(quotation_id: str, output_path: str):
    """
    Generates a pure Python PDF quotation using ReportLab.
    Matches the business style of the template.
    """
    supabase = get_supabase_client()
    
    # Fetch details
    quote = supabase.table("quotations").select("*, clients(*)").eq("id", quotation_id).execute().data[0]
    items = supabase.table("quotation_items").select("*").eq("quotation_id", quotation_id).order("item_order").execute().data
    
    # Document setup
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CompanyTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1, # Center
        textColor=colors.HexColor('#1E3D59') # Sleek dark blue
    )
    subtitle_style = ParagraphStyle(
        'CompanySub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        alignment=1,
        textColor=colors.HexColor('#17B890') # Accent green
    )
    meta_style = ParagraphStyle(
        'MetaText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=15,
        textColor=colors.HexColor('#333333')
    )
    quote_num_style = ParagraphStyle(
        'QuoteNumber',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        alignment=1,
        textColor=colors.HexColor('#17B890')
    )
    
    # Header
    story.append(Paragraph("GNB SOLUCIONES INDUSTRIALES S.A.C", title_style))
    story.append(Paragraph("Ruc: 20602667724", subtitle_style))
    story.append(Spacer(1, 15))
    
    # Date & Number
    q_date = datetime.datetime.strptime(quote["quotation_date"], "%Y-%m-%d")
    date_str = f"Callao, {q_date.day:02d} de {get_month_name(q_date.month)} del {q_date.year}"
    story.append(Paragraph(date_str, meta_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"COTIZACION N° {quote['quotation_number']}", quote_num_style))
    story.append(Spacer(1, 15))
    
    # Client block
    client_name = quote["clients"]["business_name"]
    client_attention = quote["attention_to"] or ""
    client_html = f"<b>SEÑORES:</b> {client_name}<br/><b>ATENCION:</b> {client_attention}"
    story.append(Paragraph(client_html, meta_style))
    story.append(Spacer(1, 15))
    
    # Table of items
    # Column headings
    data = [["ITEM", "SERVICIO", "CANT. UND.", "PRECIO", "TOTAL"]]
    
    cell_style = ParagraphStyle(
        'Cell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12
    )
    
    for item in items:
        # ReportLab tables don't easily support vertical cell merging out of the box
        # but we can format the service description paragraph with wordwrap and it wraps automatically
        desc_p = Paragraph(item["service_description"].replace("\n", "<br/>"), cell_style)
        
        currency_sym = "S/" if quote["currency"] == "soles" else "$"
        data.append([
            str(item["item_order"]),
            desc_p,
            f"{int(item['quantity'])} {item['unit']}",
            f"{currency_sym} {float(item['unit_price']):,.2f}",
            f"{currency_sym} {float(item['total']):,.2f}"
        ])
        
    # Table formatting
    t = Table(data, colWidths=[40, 240, 70, 70, 80])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3D59')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 9),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('TOPPADDING', (0,0), (-1,0), 8),
        ('ALIGN', (0,0), (-1,0), 'CENTER'),
        ('ALIGN', (2,1), (-1,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
        ('BOTTOMPADDING', (0,1), (-1,-1), 6),
        ('TOPPADDING', (0,1), (-1,-1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 15))
    
    # Conditions and Total
    currency_label = "Soles" if quote["currency"] == "soles" else "Dólares"
    igv_label = "Precio no Incluye IGV." if not quote["include_igv"] else "Precio Incluye IGV."
    total_str = f"{'S/' if quote['currency'] == 'soles' else '$'} {float(quote['total']):,.2f}"
    
    summary_html = f"<b>{igv_label}</b><br/><b>Moneda:</b> {currency_label}<br/>"
    summary_html += f"<b>A todo costo:</b> {total_str} ({quote['total_in_words']})"
    story.append(Paragraph(summary_html, meta_style))
    story.append(Spacer(1, 20))
    
    # Signature/Footer
    footer_html = "Agradecemos la Atención prestada.<br/><br/><b>Atte.</b><br/>Brandon Barrantes<br/>RPC: 983276501<br/>Email: gnbsolucionesindustriales@gmail.com"
    story.append(Paragraph(footer_html, meta_style))
    
    # Build
    # Get the extracted watermark image path
    watermark_image_path = get_watermark_image_path()
    
    def draw_watermark_cb(canvas, document):
        if watermark_image_path and os.path.exists(watermark_image_path):
            canvas.saveState()
            # Draw image spanning the whole page (Letter size: width=612, height=792)
            canvas.drawImage(watermark_image_path, 0, 0, width=letter[0], height=letter[1], mask='auto')
            canvas.restoreState()
            
    if watermark_image_path:
        try:
            doc.build(story, onFirstPage=draw_watermark_cb, onLaterPages=draw_watermark_cb)
        finally:
            # Clean up the watermark image temp file
            try:
                if os.path.exists(watermark_image_path):
                    os.remove(watermark_image_path)
            except Exception:
                pass
    else:
        doc.build(story)


# --- EXCEL PAYROLL EXPORT ---

def export_payroll_excel(payroll_period_id: str, output_path: str):
    """
    Generates a weekly payroll report in Excel, matching the reference spreadsheet layout.
    """
    supabase = get_supabase_client()
    
    period = supabase.table("payroll_periods").select("*").eq("id", payroll_period_id).execute().data[0]
    entries = supabase.table("payroll_entries").select("*").eq("payroll_period_id", payroll_period_id).execute().data
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Planilla de Pagos"
    
    # Enable grid lines
    ws.views.sheetView[0].showGridLines = True
    
    # 1. Styling Definitions
    font_title = Font(name="Calibri", size=16, bold=True, color="1E3D59")
    font_section = Font(name="Calibri", size=11, bold=True)
    font_header = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    font_data = Font(name="Calibri", size=10)
    font_total = Font(name="Calibri", size=10, bold=True)
    
    fill_header = PatternFill(start_color="1E3D59", end_color="1E3D59", fill_type="solid")
    fill_sub_header = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    double_bottom_border = Border(
        top=Side(style='thin', color='000000'),
        bottom=Side(style='double', color='000000')
    )
    
    # Title
    ws["A1"] = f"PLANILLA DE PAGOS - {period['title']}"
    ws["A1"].font = font_title
    ws.row_dimensions[1].height = 25
    
    # Table 1: Matrix of hours
    headers = ["PERSONAL", "PAGO POR DIA", "PAGO POR HORA", "TIPO DÍA", "MARTES", "MIERCOLES", "JUEVES", "VIERNES", "SABADO", "DOMINGO", "LUNES", "SUMA DIARIA"]
    
    ws.append([]) # row 2
    
    # Write Table 1 Header
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num)
        cell.value = header
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[3].height = 24
    
    current_row = 4
    
    for entry in entries:
        # Load days hours
        days_res = supabase.table("payroll_days") \
            .select("day_name, hours_worked, calculated_amount") \
            .eq("payroll_entry_id", entry["id"]) \
            .execute().data
            
        days_dict = {d["day_name"]: d["hours_worked"] for d in days_res}
        days_pay_dict = {d["day_name"]: d["calculated_amount"] for d in days_res}
        
        # Write Hours row
        # PERSONAL
        c1 = ws.cell(row=current_row, column=1, value=entry["employee_name_snapshot"])
        c1.font = font_section
        
        # RATES
        ws.cell(row=current_row, column=2, value=float(entry["daily_rate_snapshot"])).number_format = '0.00'
        ws.cell(row=current_row, column=3, value=float(entry["hourly_rate_snapshot"])).number_format = '0.0000'
        
        # TIPO DÍA label
        ws.cell(row=current_row, column=4, value="HORAS")
        
        # Hours per day
        days_list = ["martes", "miercoles", "jueves", "viernes", "sabado", "domingo", "lunes"]
        for idx, day in enumerate(days_list, 5):
            ws.cell(row=current_row, column=idx, value=float(days_dict.get(day, 0.0)))
            
        # Sum of hours
        total_hours = sum(float(days_dict.get(d, 0.0)) for d in days_list)
        ws.cell(row=current_row, column=12, value=total_hours)
        
        # Style the hours row
        for c in range(1, 13):
            cell = ws.cell(row=current_row, column=c)
            cell.font = font_data
            cell.fill = fill_sub_header
            cell.border = thin_border
            if c >= 2:
                cell.alignment = Alignment(horizontal="right")
                
        current_row += 1
        
        # Write Calculated Amount row
        # PERSONAL empty
        ws.cell(row=current_row, column=1, value="")
        # RATES empty
        ws.cell(row=current_row, column=2, value="")
        ws.cell(row=current_row, column=3, value="")
        
        # TIPO DÍA label
        ws.cell(row=current_row, column=4, value="PAGO DIA")
        
        # Pay per day
        for idx, day in enumerate(days_list, 5):
            ws.cell(row=current_row, column=idx, value=float(days_pay_dict.get(day, 0.0))).number_format = '0.00'
            
        # Gross Sum
        ws.cell(row=current_row, column=12, value=float(entry["gross_total"])).number_format = '0.00'
        
        # Style the payments row
        for c in range(1, 13):
            cell = ws.cell(row=current_row, column=c)
            cell.font = font_data
            cell.border = thin_border
            if c >= 5:
                cell.alignment = Alignment(horizontal="right")
                
        current_row += 1
        
    # Gross Total Summary Row
    ws.cell(row=current_row, column=1, value="SUMA TOTAL PLANILLA")
    ws.cell(row=current_row, column=12, value=float(period["total_gross"])).number_format = '0.00'
    for c in range(1, 13):
        cell = ws.cell(row=current_row, column=c)
        cell.font = font_total
        cell.border = double_bottom_border
        if c == 12:
            cell.alignment = Alignment(horizontal="right")
            
    # Save a gap of rows
    current_row += 3
    
    # Table 2: Adjustments, Net, Accounts
    ws.cell(row=current_row, column=1, value="RESUMEN NETO Y CUENTAS DE TRABAJADORES").font = font_section
    current_row += 1
    
    t2_headers = ["PERSONAL", "BRUTO (SUMA)", "DESCUENTOS / ADELANTOS", "TOTAL NETO", "N° CUENTA / YAPE", "METODO PAGO", "ESTADO"]
    for col_num, header in enumerate(t2_headers, 1):
        cell = ws.cell(row=current_row, column=col_num)
        cell.value = header
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = Alignment(horizontal="center", vertical="center")
        
    current_row += 1
    
    for entry in entries:
        ws.cell(row=current_row, column=1, value=entry["employee_name_snapshot"])
        ws.cell(row=current_row, column=2, value=float(entry["gross_total"])).number_format = '0.00'
        ws.cell(row=current_row, column=3, value=float(entry["adjustment_total"])).number_format = '0.00'
        ws.cell(row=current_row, column=4, value=float(entry["net_total"])).number_format = '0.00'
        ws.cell(row=current_row, column=5, value=entry["account_snapshot"])
        ws.cell(row=current_row, column=6, value=entry["payment_method_snapshot"].upper())
        ws.cell(row=current_row, column=7, value=entry["payment_status"].upper())
        
        for c in range(1, 8):
            cell = ws.cell(row=current_row, column=c)
            cell.font = font_data
            cell.border = thin_border
            if c in (2, 3, 4):
                cell.alignment = Alignment(horizontal="right")
                
        current_row += 1
        
    # Table 2 Total Row
    ws.cell(row=current_row, column=1, value="TOTAL GENERAL")
    ws.cell(row=current_row, column=2, value=float(period["total_gross"])).number_format = '0.00'
    ws.cell(row=current_row, column=3, value=float(period["total_adjustments"])).number_format = '0.00'
    ws.cell(row=current_row, column=4, value=float(period["total_net"])).number_format = '0.00'
    
    for c in range(1, 8):
        cell = ws.cell(row=current_row, column=c)
        cell.font = font_total
        cell.border = double_bottom_border
        if c in (2, 3, 4):
            cell.alignment = Alignment(horizontal="right")
            
    # Auto-adjust column widths
    for col in ws.columns:
        max_len = 0
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        for cell in col:
            val = str(cell.value or '')
            if len(val) > max_len:
                max_len = len(val)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)
        
    wb.save(output_path)


# --- PDF PAYROLL EXPORT ---

def export_payroll_pdf(payroll_period_id: str, output_path: str):
    """
    Generates a PDF summary of the weekly payroll.
    """
    supabase = get_supabase_client()
    
    period = supabase.table("payroll_periods").select("*").eq("id", payroll_period_id).execute().data[0]
    entries = supabase.table("payroll_entries").select("*").eq("payroll_period_id", payroll_period_id).execute().data
    
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'PayrollTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#1E3D59'),
        alignment=1
    )
    meta_style = ParagraphStyle(
        'PayrollMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=14,
        textColor=colors.HexColor('#333333')
    )
    
    story.append(Paragraph(f"REPORTE DE PLANILLA SEMANAL", title_style))
    story.append(Paragraph(f"<b>Período:</b> {period['title']}<br/><b>Día de Pago:</b> {period['payment_date']}<br/><b>Estado:</b> {period['status'].upper()}", meta_style))
    story.append(Spacer(1, 15))
    
    # Table data
    data = [["PERSONAL", "TIPO", "TARIFA DIA", "BRUTO", "AJUSTES", "NETO", "FORMA PAGO", "N° CUENTA / TELEFONO"]]
    
    for entry in entries:
        data.append([
            entry["employee_name_snapshot"],
            entry["worker_type_snapshot"].upper(),
            f"S/ {float(entry['daily_rate_snapshot']):,.2f}",
            f"S/ {float(entry['gross_total']):,.2f}",
            f"S/ {float(entry['adjustment_total']):,.2f}",
            f"S/ {float(entry['net_total']):,.2f}",
            entry["payment_method_snapshot"].upper(),
            entry["account_snapshot"] or ""
        ])
        
    # Total row
    data.append([
        "TOTAL GENERAL",
        "",
        "",
        f"S/ {float(period['total_gross']):,.2f}",
        f"S/ {float(period['total_adjustments']):,.2f}",
        f"S/ {float(period['total_net']):,.2f}",
        "",
        ""
    ])
    
    # Table layout
    t = Table(data, colWidths=[110, 50, 60, 60, 60, 60, 60, 80])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3D59')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 8),
        ('ALIGN', (0,0), (-1,0), 'CENTER'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
        ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#F2F2F2')),
        ('ALIGN', (2,1), (5,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('FONTSIZE', (0,1), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
    ]))
    
    story.append(t)
    doc.build(story)


# --- GENERAL HELPERS ---

def get_month_name(month: int) -> str:
    months = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }
    return months.get(month, "")
